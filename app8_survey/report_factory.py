import io
from math import floor

from django.conf import settings
from django.db.models import Count
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

from django.utils.translation import gettext_lazy as _

from app4_ehpad_base.models import ProfileSerenicia
from app8_survey.models import SurveyResponse, Notation, Question, Comment


def add_centered_image(document, path, height=None, width=None):
    document.add_picture(path, height=height, width=width)
    picture = document.paragraphs[-1]
    picture.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_title(document, text, level):
    title = document.add_heading(text, level=level)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def blank_lines(document, number):
    for line in range(number):
        document.add_paragraph()


def first_page(document, survey):
    add_centered_image(document, settings.IMG_LOGO, height=Cm(2.5))
    add_centered_image(document, settings.STATIC_ROOT + "/app4_ehpad_base/img/serenicia.png", height=Cm(1.5))
    blank_lines(document, 4)
    add_title(document, f'{_("Results")}\n{survey.title}\n{_(timezone.now().strftime("%B"))} {survey.year}', 0)
    document.add_page_break()


def percentage(part, whole):
    result = 100 * float(part)/float(whole)
    return f'{result:.1f}%'


def return_rate(document, survey, answers):
    add_title(document, _("Return rate"), 1)
    list_pj = ProfileSerenicia.objects.filter(user__is_active=True)
    target = {'family': list_pj.filter(user_list__user__is_active=True,
                                       user__groups__permissions__codename='view_family'),
              'resident': list_pj.filter(user__groups__permissions__codename='view_residentehpad'),
              'employees': list_pj.filter(user__groups__permissions__codename='view_internalemployees')}
    rep_count = answers.count()
    count = {'percent': percentage(rep_count, target[survey.target].count()), 'number': str(rep_count)}
    table = document.add_table(rows=2, cols=2, style='Table Grid')
    table.columns[0].width, table.columns[1].width = Cm(2.5), Cm(2.5)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head = table.rows[0].cells
    head[0].text = str(survey.year)
    head[0].merge(head[1])
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = _('Number')
    hdr_cells[1].text = _('Rate')
    row_cells = table.add_row().cells
    row_cells[0].text = count['number']
    row_cells[1].text = count['percent']


def global_percentage(question, chapter, survey):
    notations = Notation.objects.filter(question=question, chapter=chapter, chapter__survey=survey)
    if not notations:
        return '0.0%'
    choices = question.notation_choices.exclude(value=0)
    switch = floor(choices.count() / 2) + choices.first().value
    return percentage(notations.filter(notation__value__gte=switch).count(), notations.count())


def global_satisfaction(document, survey):
    add_title(document, _("Overall satisfaction rate"), 1)
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = _('Questions')
    hdr_cells[1].text = str(survey.year)
    for chapter in survey.chapters.all():
        chap_cells = table.add_row().cells
        chap_cells[0].text = chapter.title
        chap_cells[0].merge(chap_cells[1])
        for question in chapter.questions.all():
            quest_cells = table.add_row().cells
            quest_cells[0].text = question.text
            quest_cells[1].text = global_percentage(question, chapter, survey)


def result_table(document, notations, survey):
    add_title(document, _("Results"), 1)
    cols = Question.objects.filter(chapter__survey=survey).annotate(count=Count('notation_choices')).order_by('-count').first().notation_choices.count()
    table = document.add_table(rows=1, cols=cols + 1, style='Table Grid')
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for index, cell in enumerate(table.rows[0].cells):
        if index == 0:
            cell.text = _('Questions')
        else:
            cell.text = str(index)
    for chapter in survey.chapters.all():
        chap_cells = table.add_row().cells
        chap_cells[0].text = chapter.title
        for question in chapter.questions.all():
            total_notations = notations.filter(chapter=chapter, question=question)
            for index, cell in enumerate(table.add_row().cells):
                if index == 0:
                    cell.text = question.text
                else:
                    cell.text = notations.filter(notation__value=index).first().notation.text
            for index, cell in enumerate(table.add_row().cells):
                if index > 0:
                    cell.text = percentage(total_notations.filter(notation__value=index).count(),
                                           total_notations.count())


def comments(document, survey):
    add_title(document, _("Comments"), 1)
    comments_q = Comment.objects.filter(surveyresponse__survey=survey)
    for chapter in survey.chapters.all():
        add_title(document, f'{chapter.number} {chapter.title}\n', 2)
        p = document.add_paragraph()
        for com in comments_q.filter(chapter=chapter):
            p.add_run(f'«{com.text}»\n')


def build_report(survey):
    document = Document()
    answers = SurveyResponse.objects.filter(survey=survey, filling_date__isnull=False)
    notations = Notation.objects.filter(chapter__survey=survey)
    first_page(document, survey)
    return_rate(document, survey, answers)
    global_satisfaction(document, survey)
    result_table(document, notations, survey)
    comments(document, survey)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return {'file': buffer, 'name': f'{survey.title.replace(" ", "_").lower()}_{survey.year}.docx'}
