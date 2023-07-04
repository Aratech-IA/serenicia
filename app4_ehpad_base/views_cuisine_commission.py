import calendar
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from django.contrib.auth.decorators import login_required, permission_required
from django.http import FileResponse
from django.shortcuts import render
from django.templatetags.static import static
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_BUILTIN_STYLE
from docx.shared import Inches, RGBColor
import matplotlib.dates as mdates

from app4_ehpad_base.models import MenuEvaluation
from django.utils.translation import gettext_lazy as _

from app4_ehpad_base.views_cuisine_eval import get_last_day_of_month, get_evaluation
from projet.settings import settings

dark_green = RGBColor(0x00, 0xb6, 0x00)


def get_date_back_a_month(date_value):
    if date_value.month - 1 < 1:
        month = 12
        year = date_value.year - 1
    else:
        month = date_value.month - 1
        year = date_value.year
    last_day = get_last_day_of_month(datetime.strptime('1' + '/' + str(month) + '/' + str(year), '%d/%m/%Y'))
    if last_day < date_value.day:
        day = last_day
    else:
        day = date_value.day
    result = datetime.strptime(str(day) + '/' + str(month) + '/' + str(year), '%d/%m/%Y')
    return result


@login_required
@permission_required('app0_access.view_cuisineprice')
def prepare_menu_commission(request):
    if request.method == 'POST':
        start = datetime.strptime(request.POST.get('start'), '%d/%m/%Y').date()
        end = datetime.strptime(request.POST.get('end'), '%d/%m/%Y').date()
        file = get_commission_document(start, end)
        return FileResponse(file['data'], as_attachment=True, filename=file['name'])
    today = datetime.now().date()
    return render(request, 'app4_ehpad_base/cuisine_menu_commission.html',
                  {'today': today.strftime('%d/%m/%Y'), 'start': get_date_back_a_month(today).strftime('%d/%m/%Y')})


def build_evaluation_count_graph(date_value):
    month = date_value.month
    year = date_value.year - 1
    x_axis, y_axis1, y_axis2 = [], [], []
    for loop in range(0, 13):
        y_axis1.append(MenuEvaluation.objects.filter(menu__date__month=month,
                                                     menu__date__year=year, menu__type='noon').count())
        y_axis2.append(MenuEvaluation.objects.filter(menu__date__month=month,
                                                     menu__date__year=year, menu__type='evening').count())
        x_axis.append(calendar.month_abbr[month])
        month = month + 1
        if month > 12:
            month = 1
            year = year + 1
    fig, ax = plt.subplots(figsize=(8, 8))
    ind = np.arange(len(x_axis))
    width = 0.9
    p1 = ax.bar(ind, y_axis1, width, label=str(_('Noon')), color='#93a9d2')
    p2 = ax.bar(ind, y_axis2, width, bottom=y_axis1, label=str(_('Evening')), color='#5285e3')
    ax.axhline(0, color='grey', linewidth=0.8)
    ax.set_ylabel(_('Evaluations'))
    ax.set_title(_('Total evaluations per month') + ' (' + date_value.strftime('%B') + '/' + str(year - 1)
                 + ' - ' + date_value.strftime('%B/%Y') + ' )')
    ax.set_xticks(ind)
    ax.set_xticklabels(x_axis)
    ax.legend()
    ax.bar_label(p1, label_type='center')
    ax.bar_label(p2, label_type='center')
    ax.bar_label(p2)
    buffer = BytesIO()
    plt.savefig(buffer)
    buffer.seek(0)
    return buffer


def round_total(evaluation):
    total = (evaluation['dessert'] + evaluation['main_dish'] + evaluation['dessert']) / 3
    return round(total, 1)


def get_evaluation_by_type(list_eval, list_meal, meal_type):
    result = []
    for evaluation in list_meal:
        tmp_meal = None
        tmp_eval = {}
        list_recurrence = list_eval.distinct('menu__date', 'menu__type')
        if meal_type == 'entry':
            tmp_meal = evaluation.menu.entree
            tmp_eval = get_evaluation(list_eval.filter(menu__entree=tmp_meal))
            tmp_eval = {'notation': tmp_eval['notation']['entry'], 'voter': tmp_eval['voter']['entry'],
                        'recurrence': len(list_recurrence.filter(menu__entree=tmp_meal))}
        elif meal_type == 'main':
            tmp_meal = evaluation.menu.main_dish
            tmp_eval = get_evaluation(list_eval.filter(menu__main_dish=tmp_meal))
            tmp_eval = {'notation': tmp_eval['notation']['main_dish'], 'voter': tmp_eval['voter']['main_dish'],
                        'recurrence': len(list_recurrence.filter(menu__main_dish=tmp_meal))}
        elif meal_type == 'dessert':
            tmp_meal = evaluation.menu.dessert
            tmp_eval = get_evaluation(list_eval.filter(menu__dessert=tmp_meal))
            tmp_eval = {'notation': tmp_eval['notation']['dessert'], 'voter': tmp_eval['voter']['dessert'],
                        'recurrence': len(list_recurrence.filter(menu__dessert=tmp_meal))}
        tmp_eval['meal'] = tmp_meal
        result.append(tmp_eval)
    return sorted(result, key=lambda x: x['notation'])


def add_table_worst_best_to_doc(document, list_meal):
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    table.columns[0].width, table.columns[1].width, table.columns[2].width = Inches(1.2), Inches(1.2), Inches(1.2)
    table.columns[3].width = Inches(3)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = _('Notation')
    hdr_cells[1].text = _('Voter')
    hdr_cells[2].text = _('Recurrence')
    hdr_cells[3].text = _('Meal')
    for meal in list_meal:
        row_cells = table.add_row().cells
        row_cells[0].text = str(meal['notation'])
        row_cells[1].text = str(meal['voter'])
        row_cells[2].text = str(meal['recurrence'])
        try:
            row_cells[3].text = meal['meal'].name
        except AttributeError:
            row_cells[3].text = _('Not documented')
    document.add_paragraph('')


def add_table_worst_best(document, worst_list, len_table):
    para_title = _('Best')
    document.add_paragraph(f"{len_table} {para_title}")
    add_table_worst_best_to_doc(document, worst_list[::-1][:len_table])
    para_title = _('Less good')
    document.add_paragraph(f"{len_table} {para_title}")
    add_table_worst_best_to_doc(document, worst_list[:len_table])
    document.add_page_break()


def add_title(document, text, blank_space=False):
    title = document.add_heading(text, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.color.rgb = dark_green
    if blank_space:
        document.add_paragraph('')


def add_best_worst_meal_table(document, list_eval):
    worst_entry = get_evaluation_by_type(list_eval, list_eval.distinct('menu__entree'), 'entry')
    worst_main = get_evaluation_by_type(list_eval, list_eval.distinct('menu__main_dish'), 'main')
    worst_dessert = get_evaluation_by_type(list_eval, list_eval.distinct('menu__dessert'), 'dessert')
    add_title(document, _('Entry\'s notation'), True)
    add_table_worst_best(document, worst_entry, 10)
    add_title(document, _('Main course\'s notation'), True)
    add_table_worst_best(document, worst_main, 10)
    add_title(document, _('Dessert\'s notation'), True)
    add_table_worst_best(document, worst_dessert, 10)


def watzefuck(start, end):
    list_eval = MenuEvaluation.objects.prefetch_related('menu',
                                                        'menu__entree',
                                                        'menu__main_dish',
                                                        'menu__accompaniment',
                                                        'menu__dessert').filter(menu__date__gte=start,
                                                                                menu__date__lte=end)
    list_result_menu = []
    for tmp_id in list_eval.distinct('menu').values('menu'):
        tmp_data = {'menu': list_eval.distinct('menu').filter(menu__id=tmp_id['menu']).__getitem__(0).menu,
                    'evaluations': get_evaluation(list_eval.filter(menu__id=tmp_id['menu']))}
        list_result_menu.append(tmp_data)
    evaluation_selected_range = get_evaluation(list_eval)
    month_notation = round_total(evaluation_selected_range['notation'])


def build_average_notation_graph(list_eval, start, end):
    notation_range = [1, 2, 3, 4]
    list_voter, list_notation, list_days = [], [], []
    checked_day = start
    while checked_day <= end:
        tmp_eval = get_evaluation(list_eval.filter(menu__date=checked_day))
        list_notation.append(round_total(tmp_eval['notation']))
        list_voter.append(round_total(tmp_eval['voter']))
        list_days.append(checked_day)
        checked_day = checked_day + timedelta(days=1)
    fig, [ax1, ax2] = plt.subplots(2, 1, sharex=True)
    ax1.plot(list_days, list_notation, color='#5285e3')
    graph_title = _("Average notation")
    ax1.set(ylabel=_('Notation'), title=f'{graph_title} {start} / {end}')
    ax1.set_yticks(notation_range)
    ax1.set_yticklabels(notation_range)
    ax1.grid()

    locator = mdates.AutoDateLocator()
    formatter = mdates.ConciseDateFormatter(locator)
    ax2.xaxis.set_major_locator(locator)
    ax2.xaxis.set_major_formatter(formatter)
    ax2.plot(list_days, list_voter, color='#5285e3')
    graph_title = _('Average voter')
    ax2.set(ylabel=_('Voter'), title=f'{graph_title} {start} / {end}')
    ax2.set_xlim(start, end)
    ax2.grid()
    buffer = BytesIO()
    plt.savefig(buffer)
    buffer.seek(0)
    return buffer


def center_last_picture(document):
    picture = document.paragraphs[-1]
    picture.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_paragraph_title(paragraph, text):
    text = paragraph.add_run(text)
    text.underline = True
    text.bold = True


def add_cover_page(document, date_value):
    document.add_picture(settings.STATIC_ROOT + "/app4_ehpad_base/img/logo_eval-menu_80x32.png",
                         width=Inches(2.24), height=Inches(1.17))
    center_last_picture(document)
    doc_title = _('Menu commission from')
    add_title(document, f"{doc_title}    {date_value.strftime('%B %Y')}")
    document.add_paragraph()
    document.add_paragraph()
    p = document.add_paragraph()
    start_txt = _('Meeting start')
    end_txt = _('Meeting end')
    p.add_run(f"{start_txt} :\n{end_txt} :")
    para_title = _('Attendees')
    add_paragraph_title(document.add_paragraph(), f"{para_title} :")
    para_title = _('Comments')
    add_paragraph_title(document.add_paragraph(), f"{para_title} :")
    document.add_page_break()


def add_table_incoherence(document, list_meal):
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    table.columns[0].width, table.columns[1].width, table.columns[2].width = Inches(1), Inches(1), Inches(3)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = _('Notation')
    hdr_cells[1].text = _('Price')
    hdr_cells[2].text = _('Menu')
    for meal in list_meal:
        row_cells = table.add_row().cells
        row_cells[0].text = str(meal['notation'])
        row_cells[1].text = meal['price']
        row_cells[2].text = meal['menu']
    document.add_paragraph('')


def add_table_by_type(document, list_menu):
    para_title = _('Best note with lower price')
    document.add_paragraph(f"{10} {para_title}")
    add_table_incoherence(document, list_menu[:10])
    para_title = _('Less good note with higher price')
    document.add_paragraph(f"{10} {para_title}")
    add_table_incoherence(document, list_menu[::-1][:10])
    document.add_page_break()


def update_price_field(data):
    data['price'] = '{:,.2f}€'.format(data['price'] / 100)


def sort_list_incoherence(list_meal):
    result = []
    for meal in list_meal:
        tmp_data = {'notation': meal['notation']}
        try:
            tmp_data['price']: meal['meal'].price_cents
            tmp_data['menu'] = meal['meal'].name
        except AttributeError:
            tmp_data['menu'] = _('Not documented')
            tmp_data['menu'] = 0
    result = sorted(result, key=lambda x: (-x['notation'], x['price']))
    [update_price_field(tmp_data) for tmp_data in result]
    return result


def build_incoherence_table(document):
    list_eval_total = MenuEvaluation.objects.prefetch_related('menu', 'menu__entree', 'menu__main_dish',
                                                              'menu__dessert')
    eval_entry = get_evaluation_by_type(list_eval_total, list_eval_total.distinct('menu__entree'), 'entry')
    eval_main = get_evaluation_by_type(list_eval_total, list_eval_total.distinct('menu__main_dish'), 'main')
    eval_dessert = get_evaluation_by_type(list_eval_total, list_eval_total.distinct('menu__dessert'), 'entry')
    add_title(document, _('Incoherence table - Entry'), True)
    add_table_by_type(document, sort_list_incoherence(eval_entry))
    add_title(document, _('Incoherence table - Main dish'), True)
    add_table_by_type(document, sort_list_incoherence(eval_main))
    add_title(document, _('Incoherence table - Dessert'), True)
    add_table_by_type(document, sort_list_incoherence(eval_dessert))


def get_commission_document(start, end):
    graph_width = Inches(5.73)
    graph_height = Inches(4.98)
    list_eval = MenuEvaluation.objects.prefetch_related('menu',
                                                        'menu__entree',
                                                        'menu__main_dish',
                                                        'menu__accompaniment',
                                                        'menu__dessert').filter(menu__date__gte=start,
                                                                                menu__date__lte=end)
    document = Document()
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)
    logo_run = paragraph.add_run()
    logo_run.add_picture(settings.STATIC_ROOT + "/app4_ehpad_base/img/" + settings.IMG_LOGO_NAME,
                         width=Inches(8.25))
    add_cover_page(document, end)
    document.add_picture(build_evaluation_count_graph(end),
                         width=graph_width, height=graph_height)
    center_last_picture(document)
    document.add_picture(build_average_notation_graph(list_eval, start, end),
                         width=graph_width / 1.5, height=graph_height / 1.5)
    paragraph.style = document.styles['Footer']
    center_last_picture(document)
    document.add_page_break()
    add_best_worst_meal_table(document, list_eval)
    build_incoherence_table(document)
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.gutter = Inches(0)
    for paragraph in document.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return {'data': buffer, 'name': f"commission{end.strftime('%B%Y')}.docx"}
